import pdfplumber
import spacy
import re
from collections import OrderedDict
from docx import Document

# Load spaCy NLP model
nlp = spacy.load("en_core_web_sm")

# Define common lists (can be extended)
COMMON_SKILLS = [
    'python', 'java', 'c++', 'machine learning', 'deep learning',
    'data analysis', 'sql', 'html', 'css'
]
COMMON_LANGUAGES = [
    'english', 'hindi', 'french', 'german', 'spanish', 'mandarin'
]

def clean_line(line):
    """Remove leading bullets, dashes, and whitespace."""
    return line.strip().lstrip("-• ").strip()

def extract_text(file):
    """Extract raw text from a PDF file using pdfplumber."""
    text = ''
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
    return text.lower()

def extract_name(text):
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return None

def extract_email(text):
    match = re.search(r'[\w\.-]+@[\w\.-]+', text)
    return match.group(0) if match else None

def extract_phone(text):
    match = re.search(r'(\+?\d{1,3}[\s-]?)?(\(?\d{2,4}\)?[\s-]?)?\d{3,4}[\s-]?\d{3,4}', text)
    return match.group(0) if match else None


def extract_education(text):
    education_keywords = [
        'bachelor', 'master', 'phd', 'university', 'college',
        'b.tech', 'm.tech', 'bsc', 'msc'
    ]
    lines = text.split('\n')
    education_lines = [
        clean_line(line) for line in lines
        if any(keyword in line.lower() for keyword in education_keywords)
    ]
    return education_lines

def extract_skills(text):
    return [skill for skill in COMMON_SKILLS if skill in text]

def extract_experience(text):
    experience_keywords = [
        'experience', 'internship', 'worked at', 'company',
        'role', 'position'
    ]
    lines = text.split('\n')
    experience_lines = []
    for line in lines:
        if any(keyword in line.lower() for keyword in experience_keywords):
            cleaned = clean_line(line)
            if cleaned.lower() not in experience_keywords:
                experience_lines.append(cleaned)
    return experience_lines

def extract_certifications(text):
    cert_keywords = ['certification', 'certificate', 'certified', 'certifications']
    lines = text.split('\n')
    cert_lines = []

    for line in lines:
        line_lower = line.lower().strip()
        if any(keyword in line_lower for keyword in cert_keywords):
            cleaned = clean_line(line)
            # ❗ Skip if line is just a keyword or heading
            if cleaned.lower() not in cert_keywords and len(cleaned.split()) > 2:
                cert_lines.append(cleaned)

    return cert_lines


def extract_languages(text):
    return [lang for lang in COMMON_LANGUAGES if lang in text]

def parse_resume(text):
    return OrderedDict([
        ("Name", extract_name(text)),
        ("Email", extract_email(text)),
        ("Phone", extract_phone(text)),
        ("Education", extract_education(text)),
        ("Skills", extract_skills(text)),
        ("Experience", extract_experience(text)),
        ("Certifications", extract_certifications(text)),
        ("Languages", extract_languages(text))
    ])


def save_to_word(parsed_data, output_path="parsed_resume.docx"):
    """Save parsed data into a .docx Word document."""
    doc = Document()
    doc.add_heading("Parsed Resume Details", level=1)

    for section, value in parsed_data.items():
        doc.add_heading(section, level=2)
        if isinstance(value, list):
            for item in value:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph(str(value))

    doc.save(output_path)
    return output_path
